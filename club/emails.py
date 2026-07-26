"""Email notifications for the contact form and tournament registrations.

Sends via whatever EMAIL_BACKEND is configured (see settings.py) — SMTP if
EMAIL_HOST_USER/EMAIL_HOST_PASSWORD are set, otherwise the console backend,
which just prints the email to the terminal running `runserver` instead of
actually sending it (useful for local development without real credentials).

Failures are logged loudly (not swallowed) so a wrong password / SMTP block
/ etc. shows up clearly in the terminal instead of silently doing nothing —
but they never raise, so a broken mail server never breaks the page for the
visitor submitting the form; the submission is always saved to the database
and visible in the admin regardless of whether the email goes out.
"""

import io
import logging

from django.conf import settings
from django.core.mail import EmailMessage

logger = logging.getLogger(__name__)


def _notification_recipient(club):
    return (
        (club and club.email)
        or settings.CONTACT_NOTIFICATION_EMAIL
        or settings.DEFAULT_FROM_EMAIL
    )


def _send(subject, body, recipient=None, reply_to=None, attachments=None, recipients=None):
    to = list(recipients or [])
    if recipient:
        to.append(recipient)
    # De-dupe while preserving order.
    seen = set()
    to = [addr for addr in to if addr and not (addr in seen or seen.add(addr))]
    if not to:
        message = (
            "Email not sent: no recipient configured. Set an email address "
            "on Club Info in the admin, or set CONTACT_NOTIFICATION_EMAIL."
        )
        logger.warning(message)
        print(f"[club.emails] {message}")
        return

    email = EmailMessage(
        subject=subject,
        body=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=to,
        reply_to=[reply_to] if reply_to else None,
    )
    for filename, content, mimetype in attachments or []:
        email.attach(filename, content, mimetype)

    try:
        email.send(fail_silently=False)
    except Exception as exc:  # noqa: BLE001 - we want to log *any* send failure
        message = (
            f"Failed to send email to {to} using "
            f"{settings.EMAIL_BACKEND} (host={settings.EMAIL_HOST}, "
            f"user={settings.EMAIL_HOST_USER or '(not set)'}): {exc!r}\n"
            "If you haven't already, add EMAIL_HOST_USER and "
            "EMAIL_HOST_PASSWORD to your .env file (see .env.example) and "
            "restart the server. For Gmail you need an App Password, not "
            "your normal password."
        )
        logger.error(message)
        print(f"[club.emails] {message}")
    else:
        print(f"[club.emails] Email sent to {to}: {subject}")


def _roster_rows(registration):
    """Returns (jersey_number_str, name) pairs for named players only."""
    named = [p for p in registration.players.all() if p.name.strip()]
    return [
        (str(p.jersey_number) if p.jersey_number is not None else "-", p.name)
        for p in named
    ]


def build_registration_docx(registration, club):
    """Builds a Word (.docx) document summarising a tournament registration,
    returned as raw bytes ready to attach to an email.
    """
    from docx import Document
    from docx.shared import Pt

    document = Document()

    title = document.add_heading("Tournament Team Registration", level=1)
    title.runs[0].font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(
        f"{club.name if club else 'Gurkhali FC'} — {registration.tournament_name}"
    )
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(13)

    document.add_paragraph(
        f"Submitted: {registration.submitted_at.strftime('%d %B %Y, %I:%M %p')}"
    ).italic = True

    document.add_heading("Team Details", level=2)
    _add_docx_table(
        document,
        [
            ("Tournament", registration.tournament_name),
            ("Team Name", registration.team_name),
            ("Division", registration.get_division_display()),
            ("Number of Players", str(registration.player_count)),
            ("Home City / Suburb", registration.home_city or "-"),
            ("Status", registration.get_status_display()),
        ],
    )

    document.add_heading("Contact", level=2)
    _add_docx_table(
        document,
        [
            ("Manager / Coach", registration.manager_name),
            ("Phone", registration.phone),
            ("Email", registration.email),
        ],
    )

    roster = _roster_rows(registration)
    document.add_heading("Player Roster", level=2)
    if roster:
        table = document.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        header[0].text = "#"
        header[1].text = "Player Name"
        for jersey, name in roster:
            row = table.add_row().cells
            row[0].text = jersey
            row[1].text = name
    else:
        document.add_paragraph("No players listed.")

    document.add_heading("Additional Information", level=2)
    document.add_paragraph("Previous tournament experience:").bold = True
    document.add_paragraph(registration.experience or "None provided.")
    document.add_paragraph("Notes:").bold = True
    document.add_paragraph(registration.notes or "None provided.")

    document.add_paragraph()
    agreement = document.add_paragraph()
    agreement.add_run(
        "Confirmed agreement to tournament rules and code of conduct."
        if registration.agreed_to_rules
        else "Did NOT confirm agreement to tournament rules."
    ).bold = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_docx_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    return table


def build_registration_pdf(registration, club):
    """Builds a PDF summarising a tournament registration, mirroring the
    Word doc above, returned as raw bytes ready to attach to an email.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, "Tournament Team Registration", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(179, 18, 42)
    pdf.cell(
        0,
        8,
        f"{club.name if club else 'Gurkhali FC'} - {registration.tournament_name}",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.set_text_color(0, 0, 0)

    pdf.set_font("Helvetica", "I", 10)
    pdf.cell(
        0,
        8,
        f"Submitted: {registration.submitted_at.strftime('%d %B %Y, %I:%M %p')}",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(2)

    _pdf_section(pdf, "Team Details")
    _pdf_table(
        pdf,
        [
            ("Tournament", registration.tournament_name),
            ("Team Name", registration.team_name),
            ("Division", registration.get_division_display()),
            ("Number of Players", str(registration.player_count)),
            ("Home City / Suburb", registration.home_city or "-"),
            ("Status", registration.get_status_display()),
        ],
    )

    _pdf_section(pdf, "Contact")
    _pdf_table(
        pdf,
        [
            ("Manager / Coach", registration.manager_name),
            ("Phone", registration.phone),
            ("Email", registration.email),
        ],
    )

    _pdf_section(pdf, "Player Roster")
    roster = _roster_rows(registration)
    if roster:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(20, 8, "#", border=1, fill=True)
        pdf.cell(0, 8, "Player Name", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for jersey, name in roster:
            pdf.cell(20, 8, jersey, border=1)
            pdf.cell(0, 8, name, border=1, new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, "No players listed.", new_x="LMARGIN", new_y="NEXT")

    _pdf_section(pdf, "Additional Information")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 7, "Previous tournament experience:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, registration.experience or "None provided.")
    pdf.ln(1)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 7, "Notes:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, registration.notes or "None provided.")

    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 10)
    pdf.multi_cell(
        0,
        6,
        "Confirmed agreement to tournament rules and code of conduct."
        if registration.agreed_to_rules
        else "Did NOT confirm agreement to tournament rules.",
    )

    return bytes(pdf.output())


def _pdf_section(pdf, title):
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(11, 44, 96)
    pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)


def _pdf_table(pdf, rows):
    pdf.set_font("Helvetica", "B", 10)
    for label, value in rows:
        pdf.cell(55, 7, label, border=1)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, str(value), border=1, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "B", 10)


def send_contact_notification(contact_message, club):
    _send(
        subject=f"[{club.name if club else 'Gurkhali FC'}] New contact message from {contact_message.name}",
        body=(
            f"You have a new message from the club website contact form.\n\n"
            f"Name: {contact_message.name}\n"
            f"Email: {contact_message.email}\n\n"
            f"Message:\n{contact_message.message}\n\n"
            f"Reply directly to {contact_message.email}, or view/manage this "
            f"message in the Django admin under Contact Messages."
        ),
        recipient=_notification_recipient(club),
        reply_to=contact_message.email,
    )


def send_registration_notification(registration, club):
    docx_bytes = build_registration_docx(registration, club)
    pdf_bytes = build_registration_pdf(registration, club)
    safe_team_name = "".join(
        c for c in registration.team_name if c.isalnum() or c in (" ", "-", "_")
    ).strip() or "team"

    roster = _roster_rows(registration)
    roster_lines = "\n".join(f"  #{jersey} — {name}" for jersey, name in roster) or "  (none listed)"
    club_name = club.name if club else "Gurkhali FC"

    _send(
        subject=f"[{club_name}] New tournament registration: {registration.team_name}",
        body=(
            f"A new team has registered for a tournament via the club website.\n\n"
            f"Tournament: {registration.tournament_name}\n"
            f"Team name: {registration.team_name}\n"
            f"Division: {registration.get_division_display()}\n"
            f"Manager/coach: {registration.manager_name}\n"
            f"Phone: {registration.phone}\n"
            f"Email: {registration.email}\n"
            f"Number of players: {registration.player_count}\n"
            f"Home city/suburb: {registration.home_city or '-'}\n"
            f"Previous tournament experience: {registration.experience or '-'}\n"
            f"Notes: {registration.notes or '-'}\n\n"
            f"Player roster:\n{roster_lines}\n\n"
            f"A Word document and a PDF with these details (including the "
            f"full roster) are attached. Manage this registration "
            f"(approve/reject) in the Django admin under Team Registrations."
        ),
        recipient=_notification_recipient(club),
        reply_to=registration.email,
        attachments=[
            (
                f"Registration - {safe_team_name}.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            (
                f"Registration - {safe_team_name}.pdf",
                pdf_bytes,
                "application/pdf",
            ),
        ],
    )


def send_registration_received_confirmation(registration, club):
    """Tell the registering team their application is under review."""
    club_name = club.name if club else "Gurkhali FC"
    _send(
        subject=f"[{club_name}] Registration received — {registration.team_name}",
        body=(
            f"Hi {registration.manager_name},\n\n"
            f"Thanks for registering {registration.team_name} for "
            f"{registration.tournament_name}.\n\n"
            "Your team registration is being reviewed and when approved you "
            "will get notified.\n\n"
            "Once all 16 teams are approved, you will receive the match "
            "schedules by email.\n\n"
            f"Division: {registration.get_division_display()}\n"
            f"Players listed: {registration.player_count}\n\n"
            f"— {club_name}"
        ),
        recipient=registration.email,
        reply_to=_notification_recipient(club),
    )


def send_schedule_ready_notifications(club):
    """Email every approved team that match schedules are ready."""
    from .models import TeamRegistration

    club_name = club.name if club else "Gurkhali FC"
    approved = list(
        TeamRegistration.objects.filter(status=TeamRegistration.Status.APPROVED)
        .exclude(email="")
        .order_by("team_name")
    )
    if not approved:
        return 0

    site_hint = "Visit the Gurkhali FC website Schedule section for fixtures and group tables."
    sent = 0
    for reg in approved:
        _send(
            subject=f"[{club_name}] Match schedules are ready — {reg.tournament_name}",
            body=(
                f"Hi {reg.manager_name},\n\n"
                f"Great news — all 16 teams for {reg.tournament_name} have "
                f"been approved.\n\n"
                f"Match schedules are now available for {reg.team_name}.\n"
                f"{site_hint}\n\n"
                f"— {club_name}"
            ),
            recipient=reg.email,
            reply_to=_notification_recipient(club),
        )
        sent += 1
    return sent
